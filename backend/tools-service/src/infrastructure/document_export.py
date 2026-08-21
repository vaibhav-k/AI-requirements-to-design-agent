"""Deterministic, LLM-free technical design document rendering to
``.docx``.

The technical-design analogue of ``work_breakdown_export.py``: pure
logic, no I/O beyond Graphviz's own diagram rendering, no Azure OpenAI
dependency - the Technical Writer agent (on the orchestrator) decides
*what* the document says; this module owns *how* that structured content
becomes an actual Word document (outline numbering, a Table of Contents
field, the embedded architecture diagram, tables, and a requirements
traceability appendix).

Adapted from - not copied verbatim from - Parnell-AI-Persona-Agent's own
``document_creation_agent/operations.py``. Two real differences, both
forced by this project's own domain shapes rather than a stylistic
choice:

* This project's ``Requirement`` has no ``title``/``category`` field (see
  ``src.domain.requirements``), so the traceability appendix table here
  has fewer columns than Parnell's.
* This project's ``WorkBreakdownArtifact`` has no sprint/story-point
  concept (see ``src.domain.work_breakdown`` - just Feature -> Story ->
  Task with an XS-XL ``EffortEstimate``), so there is no delivery-summary
  appendix here; the work breakdown is instead summarized as a simple
  feature/story/task count line.

Two details worth knowing, both carried over unchanged from Parnell's
approach since they're genuine Word/python-docx constraints, not design
choices:

* **The TOC is a Word field, not computed here.** python-docx cannot
  compute page numbers, so the document carries a ``TOC`` field plus
  ``w:updateFields`` in settings.xml, which makes Word populate it when
  the file is opened.
* **The diagram is rendered to PNG, not the SVG the UI uses.** python-docx
  cannot embed SVG, so this calls
  ``ArchitectureDiagramGenerator.generate_png`` rather than
  ``generate``.
"""

from __future__ import annotations

import base64
import io
import logging
from datetime import UTC, datetime

from docx import Document as _new_document
from docx.document import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from docx.text.paragraph import Paragraph

from src.domain.design import DiagramMetadata, SystemDesignArtifact
from src.domain.errors import DiagramGenerationError, TechnicalDesignExportError
from src.domain.requirements import RequirementsArtifact
from src.domain.technical_design import (
    DesignSection,
    DesignTable,
    TechnicalDesignExport,
    TechnicalDesignExportRequest,
)
from src.domain.work_breakdown import WorkBreakdownArtifact
from src.infrastructure.diagram import ArchitectureDiagramGenerator

logger = logging.getLogger(__name__)

MAX_LEVEL = 3
TABLE_STYLE = "Table Grid"
CONTENT_WIDTH_INCHES = 6.0


# --------------------------------------------------------------------------- #
# Word field plumbing
# --------------------------------------------------------------------------- #


def _add_field(paragraph: Paragraph, instruction: str, placeholder: str = "") -> None:
    """Insert a Word field (e.g. ``TOC \\o "1-3"``, ``PAGE``) into a
    paragraph.

    A field is five runs, one per state transition - begin, instruction
    text, separate, cached result, end - not one run holding every
    ``w:fldChar``/``w:instrText``/``w:t`` child. Cramming them into a
    single run (an earlier version of this function did exactly that) is
    the kind of malformed field Word tolerates just well enough to *open*
    without complaint, but then refuses to actually recompute: the TOC
    field kept showing its placeholder text forever, even after
    right-click > Update Field, because Word's field-state parser never
    reliably resolves the begin/separate/end boundaries when they share a
    run with the instruction text and the cached result. Splitting each
    transition into its own run, and marking ``begin`` dirty, is the
    standard OOXML recipe (also how Word itself writes fields) and is
    what actually makes the field recompute.
    """
    begin_run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    begin_run._r.append(begin)

    instr_run = paragraph.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    instr_run._r.append(instr)

    separate_run = paragraph.add_run()
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run._r.append(separate)

    if placeholder:
        paragraph.add_run(placeholder)

    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def _mark_fields_dirty(document: Document) -> None:
    """Ask Word to refresh fields when the document opens, so the TOC
    fills in."""
    settings = document.settings.element
    if settings.find(qn("w:updateFields")) is None:
        update = OxmlElement("w:updateFields")
        update.set(qn("w:val"), "true")
        settings.append(update)


# --------------------------------------------------------------------------- #
# Building blocks
# --------------------------------------------------------------------------- #


def _outline_numbers(sections: list[DesignSection]) -> list[str]:
    """Derive ``"1"``, ``"1.1"``, ``"1.1.1"`` from the level sequence.

    Same counters-per-depth walk the work breakdown exporter's WBS codes
    use; levels are clamped so a stray level still lands somewhere
    sensible.
    """
    counters: list[int] = []
    numbers: list[str] = []
    for section in sections:
        depth = max(1, min(section.level, MAX_LEVEL)) - 1
        del counters[depth + 1 :]
        if len(counters) == depth + 1:
            counters[depth] += 1
        else:
            counters.extend([1] * (depth + 1 - len(counters)))
        numbers.append(".".join(str(c) for c in counters))
    return numbers


def _add_paragraphs(document: Document, text: str) -> None:
    """Add each non-blank line of prose as its own paragraph."""
    for line in (text or "").splitlines():
        if line.strip():
            document.add_paragraph(line.strip())


def _add_table(document: Document, table: DesignTable) -> bool:
    """Render a table. Ragged rows are padded/trimmed so Word never gets
    bad XML."""
    if not table.rows:
        return False
    headers = table.headers or [f"Column {i + 1}" for i in range(len(table.rows[0]))]
    width = len(headers)

    word_table = document.add_table(rows=1, cols=width)
    word_table.style = TABLE_STYLE
    for cell, header in zip(word_table.rows[0].cells, headers, strict=False):
        cell.text = str(header)
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    for row in table.rows:
        cells = word_table.add_row().cells
        padded = list(row[:width]) + [""] * max(0, width - len(row))
        for cell, value in zip(cells, padded, strict=False):
            cell.text = "" if value is None else str(value)

    if table.caption:
        document.add_paragraph(table.caption, style="Caption")
    return True


def _diagram_metadata(
    title: str, description: str, design: SystemDesignArtifact
) -> DiagramMetadata:
    """Deterministic metadata block for an embedded diagram - see
    ``DiagramMetadata``'s docstring: never invented, "TBD" for anything
    genuinely unknown at export time (there is no design-version number
    threaded into a technical-design export request, so ``version``
    stays at its default of ``1`` rather than guessing one)."""

    return DiagramMetadata(
        title=title,
        description=description,
        scope=design.architecture_summary[:280],
        author="TBD",
        last_updated=datetime.now(UTC).isoformat(),
        external_references=[],
    )


def _diagram_pngs(
    generator: ArchitectureDiagramGenerator, request: TechnicalDesignExportRequest
) -> tuple[bytes | None, bytes | None, list[str]]:
    """Render both required architecture diagrams (Logical Architecture +
    Azure Service Mapping) for ``request.design`` as PNG, for embedding.

    Never raises - a failed diagram render degrades to an omitted figure
    plus a warning, exactly as Parnell's own ``_diagram_png`` does, since
    losing one figure shouldn't fail the whole document export. The two
    diagrams are rendered independently, so one failing doesn't take the
    other down with it.
    """
    if not request.design.components:
        return None, None, []

    warnings: list[str] = []

    try:
        logical_png = generator.generate_logical_png(
            request.design,
            _diagram_metadata(
                "Logical Architecture Diagram",
                "Technology-agnostic components, actors, and their "
                "interactions/trust boundaries.",
                request.design,
            ),
        )
    except DiagramGenerationError as exc:
        logger.warning("Logical diagram PNG render failed", exc_info=True)
        logical_png = None
        warnings.append(f"Logical architecture diagram omitted - render failed: {exc}")

    try:
        azure_png = generator.generate_azure_mapping_png(
            request.design,
            _diagram_metadata(
                "Azure Service Mapping Diagram",
                "Every major logical component mapped to its concrete "
                "Azure service implementation, plus supporting Azure "
                "services.",
                request.design,
            ),
        )
    except DiagramGenerationError as exc:
        logger.warning("Azure mapping diagram PNG render failed", exc_info=True)
        azure_png = None
        warnings.append(f"Azure service mapping diagram omitted - render failed: {exc}")

    return logical_png, azure_png, warnings


def _azure_mapping_table(design: SystemDesignArtifact) -> DesignTable | None:
    """A real, screen-reader-accessible Word table listing every
    component/actor/external-dependency -> Azure service mapping, plus
    the Azure-mapping-only supporting services - the accessibility
    requirement that trust-zone/connectivity/rationale information never
    depend on the diagram's color coding alone.
    """
    if not design.azure_mappings and not design.supporting_azure_services:
        return None

    display_name_by_id: dict[str, str] = {}
    for component in design.components:
        display_name_by_id[component.id] = component.name
    for actor in design.actors:
        display_name_by_id[actor.id] = actor.name
    for dependency in design.external_dependencies:
        display_name_by_id[dependency.id] = dependency.name

    rows: list[list[str]] = [
        [
            f"{mapping.component_id} - "
            f"{display_name_by_id.get(mapping.component_id, '?')}",
            mapping.azure_service,
            mapping.service_category,
            mapping.connectivity,
            mapping.trust_zone,
            mapping.rationale,
        ]
        for mapping in design.azure_mappings
    ]

    rows.extend(
        [
            f"(supporting) {service.id}",
            service.azure_service,
            service.category,
            "-",
            "-",
            service.rationale,
        ]
        for service in design.supporting_azure_services
    )

    return DesignTable(
        headers=[
            "Component",
            "Azure Service",
            "Category",
            "Connectivity",
            "Trust Zone",
            "Rationale",
        ],
        rows=rows,
        caption="Table 1 - Logical component to Azure service mapping",
    )


def _add_diagram(document: Document, png: bytes, caption: str) -> None:
    document.add_picture(io.BytesIO(png), width=Inches(CONTENT_WIDTH_INCHES))
    document.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    figure = document.add_paragraph(caption, style="Caption")
    figure.alignment = WD_ALIGN_PARAGRAPH.CENTER


# --------------------------------------------------------------------------- #
# Front matter / appendices
# --------------------------------------------------------------------------- #


def _add_title_page(document: Document, request: TechnicalDesignExportRequest) -> None:
    doc = request.document
    title = document.add_paragraph(doc.document_title or "Technical Design Document")
    title.style = "Title"
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    for line in (doc.system_name, f"Version {doc.version}" if doc.version else ""):
        if line:
            paragraph = document.add_paragraph(line)
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.font.size = Pt(12)
    document.add_page_break()  # type: ignore[no-untyped-call]


def _add_toc(document: Document) -> None:
    heading = document.add_paragraph("Table of Contents")
    heading.style = "Heading 1"
    _add_field(
        document.add_paragraph(),
        'TOC \\o "1-3" \\h \\z \\u',
        "Right-click and choose Update Field to build the table of contents.",
    )
    document.add_page_break()  # type: ignore[no-untyped-call]


def _add_footer(document: Document) -> None:
    footer = document.sections[0].footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.text = "Page "
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _add_field(paragraph, "PAGE", "1")


def _requirement_rows(requirements: RequirementsArtifact) -> list[list[str]]:
    """Flatten every functional/non-functional requirement into one
    traceability row each - the technical-design analogue of
    ``work_breakdown_export.py``'s own flattening, minus the
    title/category columns this project's ``Requirement`` doesn't have.
    """
    rows: list[list[str]] = []
    for requirement in (
        *requirements.functional_requirements,
        *requirements.non_functional_requirements,
    ):
        rows.append(
            [
                requirement.id,
                requirement.description,
                requirement.priority,
            ]
        )
    return rows


def _add_traceability_appendix(
    document: Document, requirements: RequirementsArtifact
) -> int:
    """Appendix A: the approved requirements, rendered from the artifact
    itself."""
    heading = document.add_paragraph("Appendix A - Requirements Traceability")
    heading.style = "Heading 1"

    rows = _requirement_rows(requirements)
    if not rows:
        document.add_paragraph("No requirements were supplied.")
        return 0

    return int(
        _add_table(
            document,
            DesignTable(
                headers=["ID", "Requirement", "Priority"],
                rows=rows,
            ),
        )
    )


def _add_work_breakdown_appendix(
    document: Document, work_breakdown: WorkBreakdownArtifact
) -> None:
    """Appendix B: a simple Feature/Story/Task count summary - this
    project's work breakdown has no sprint/story-point concept for a
    fuller delivery-plan appendix like Parnell's."""
    heading = document.add_paragraph("Appendix B - Implementation Work Summary")
    heading.style = "Heading 1"

    feature_count = len(work_breakdown.features)
    story_count = sum(len(feature.stories) for feature in work_breakdown.features)
    task_count = sum(
        len(story.tasks)
        for feature in work_breakdown.features
        for story in feature.stories
    )

    document.add_paragraph(
        f"{feature_count} feature(s), {story_count} user story(ies), and "
        f"{task_count} task(s) planned to implement this design."
    )


# --------------------------------------------------------------------------- #
# Tool entry point
# --------------------------------------------------------------------------- #


class TechnicalDesignExporter:
    """Render a ``TechnicalDesignArtifact`` to a downloadable ``.docx``.

    The ``DocumentExporterPort`` implementation's counterpart on this
    side of the design-tools split - see
    ``app.application.ports.DocumentExporterPort`` (orchestrator) for the
    port this ultimately backs, reached via
    ``backend/mcp-wrapper``'s ``export_technical_design_tool``.
    """

    def __init__(self) -> None:
        self._diagram_generator = ArchitectureDiagramGenerator()

    def export(self, request: TechnicalDesignExportRequest) -> TechnicalDesignExport:
        """Render ``request.document`` to a base64 ``.docx``.

        Raises ``TechnicalDesignExportError`` if the document has no
        sections at all - nothing meaningful to render otherwise. A
        diagram render failure degrades to an omitted figure plus a
        warning rather than failing the export - see ``_diagram_png``.
        """
        doc = request.document

        if not doc.sections:
            raise TechnicalDesignExportError(
                "The technical design document has no sections to render."
            )

        warnings: list[str] = []
        logical_png, azure_png, diagram_warnings = _diagram_pngs(
            self._diagram_generator, request
        )
        warnings.extend(diagram_warnings)

        document = _new_document()
        _mark_fields_dirty(document)
        _add_title_page(document, request)
        _add_toc(document)
        _add_footer(document)

        headings = 0
        tables = 0
        diagram_embedded = False

        if doc.executive_summary.strip():
            heading = document.add_paragraph("Executive Summary")
            heading.style = "Heading 1"
            headings += 1
            _add_paragraphs(document, doc.executive_summary)

        # Both diagrams belong to whichever section asked for one; if
        # none did, place them up front rather than dropping the
        # figures.
        anchor_index = next(
            (i for i, s in enumerate(doc.sections) if s.include_diagram), None
        )
        logical_caption = (
            f"Figure 1 - {doc.system_name or 'System'} logical architecture"
        )
        azure_caption = (
            f"Figure 2 - {doc.system_name or 'System'} Azure service mapping"
        )

        if anchor_index is None:
            if logical_png:
                _add_diagram(document, logical_png, logical_caption)
                diagram_embedded = True
            if azure_png:
                _add_diagram(document, azure_png, azure_caption)
                diagram_embedded = True

        numbers = _outline_numbers(doc.sections)
        sections_with_numbers = list(zip(doc.sections, numbers, strict=True))
        for index, (section, number) in enumerate(sections_with_numbers):
            level = max(1, min(section.level, MAX_LEVEL))
            heading = document.add_paragraph(
                f"{number}  {section.title.strip() or 'Untitled'}"
            )
            heading.style = f"Heading {level}"
            headings += 1

            _add_paragraphs(document, section.body)

            for bullet in section.bullets:
                if bullet.strip():
                    document.add_paragraph(bullet.strip(), style="List Bullet")

            for step in section.numbered_steps:
                if step.strip():
                    document.add_paragraph(step.strip(), style="List Number")

            if index == anchor_index:
                if logical_png:
                    _add_diagram(document, logical_png, logical_caption)
                    diagram_embedded = True
                if azure_png:
                    _add_diagram(document, azure_png, azure_caption)
                    diagram_embedded = True

            if section.table:
                tables += int(_add_table(document, section.table))

        if (
            request.design.components
            and not diagram_embedded
            and not logical_png
            and not azure_png
        ):
            warnings.append("The architecture diagrams could not be embedded.")

        document.add_page_break()  # type: ignore[no-untyped-call]
        tables += _add_traceability_appendix(document, request.requirements)
        headings += 1

        _add_work_breakdown_appendix(document, request.work_breakdown)

        azure_table = _azure_mapping_table(request.design)
        if azure_table is not None:
            heading = document.add_paragraph("Appendix C - Azure Service Mapping")
            heading.style = "Heading 1"
            headings += 1
            tables += int(_add_table(document, azure_table))
        headings += 1

        buffer = io.BytesIO()
        document.save(buffer)
        data = buffer.getvalue()

        return TechnicalDesignExport(
            docx_base64=base64.b64encode(data).decode("ascii"),
            filename="technical-design.docx",
            heading_count=headings,
            table_count=tables,
            diagram_embedded=diagram_embedded,
            byte_count=len(data),
            warnings=warnings,
        )
